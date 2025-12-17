"""
Text Extraction Utilities
Handles text extraction from different document formats.
"""

import io
from typing import Dict, Any, Optional
from pypdf import PdfReader
from docx import Document as DocxDocument


class TextExtractor:
    """
    Extracts text from various document formats.
    """

    def extract_text(self, file_content: bytes, format_result: Dict[str, Any]) -> str:
        """
        Extract text from document based on detected format.

        Args:
            file_content: Raw file bytes
            format_result: Format detection results from FormatDetector

        Returns:
            str: Extracted text content

        Raises:
            ValueError: If format is not supported or extraction fails
        """
        format_type = format_result.get("format_type")

        if format_type == "pdf":
            return self._extract_pdf_text(file_content)
        elif format_type == "docx":
            return self._extract_docx_text(file_content)
        elif format_type == "image":
            raise ValueError("Image processing not yet supported")
        else:
            raise ValueError(f"Unsupported format for text extraction: {format_type}")

    def _extract_pdf_text(self, file_content: bytes) -> str:
        """
        Extract text from PDF using PyMuPDF (fitz).

        Args:
            file_content: PDF file bytes

        Returns:
            str: Extracted text
        """
        try:
            import fitz  # PyMuPDF

            # Open PDF from bytes
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""

            # Extract text from all pages
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"

            doc.close()
            return text.strip()

        except ImportError:
            # Fallback to pypdf if PyMuPDF not available
            print("Warning: PyMuPDF not available, falling back to pypdf")
            return self._extract_pdf_text_pypdf(file_content)
        except Exception as e:
            raise ValueError(f"PDF text extraction failed: {str(e)}")

    def _extract_pdf_text_pypdf(self, file_content: bytes) -> str:
        """
        Fallback PDF extraction using pypdf.

        Args:
            file_content: PDF file bytes

        Returns:
            str: Extracted text
        """
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text += page_text + "\n"

            return text.strip()

        except Exception as e:
            raise ValueError(f"PDF fallback extraction failed: {str(e)}")

    def _extract_docx_text(self, file_content: bytes) -> str:
        """
        Extract text from DOCX document.

        Args:
            file_content: DOCX file bytes

        Returns:
            str: Extracted text
        """
        try:
            # Create a BytesIO object for python-docx
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            return text.strip()

        except Exception as e:
            raise ValueError(f"DOCX text extraction failed: {str(e)}")

    def get_extraction_stats(self, text: str, format_result: Dict[str, Any]) -> Dict[str, Any]:
        """
        Get statistics about the extracted text.

        Args:
            text: Extracted text content
            format_result: Format detection results

        Returns:
            dict: Extraction statistics
        """
        return {
            "format_type": format_result.get("format_type"),
            "text_length": len(text),
            "word_count": len(text.split()) if text else 0,
            "line_count": len(text.split('\n')) if text else 0,
            "has_content": bool(text.strip()),
            "is_scanned": format_result.get("is_scanned", False)
        }
